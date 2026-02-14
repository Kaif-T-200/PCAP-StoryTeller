from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
import os
import io
import uuid
import subprocess
import json
import sys
from datetime import datetime
from werkzeug.utils import secure_filename
from collections import defaultdict

app = Flask(__name__, static_folder='../static', template_folder='../templates')
ALLOWED_EXTENSIONS = {'pcap', 'pcapng', 'cap'}
# Use absolute path for upload folder to avoid cwd issues
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

try:
    import scapy
    SCAPY_OK = True
except ImportError:
    SCAPY_OK = False
    print("❌ Scapy not installed. Run: pip install scapy", file=sys.stderr)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/timeline')
def timeline():
    return render_template('timeline.html')

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/analytics')
def analytics():
    return render_template('analytics.html')

@app.route('/threats')
def threats():
    return render_template('threats.html')

@app.route('/search')
def search_page():
    return render_template('search.html')

@app.route('/geolocation')
def geolocation():
    return render_template('geolocation.html')

@app.route('/api/analytics')
def api_analytics():
    data = load_report_data()
    if not data:
        return jsonify({'error': 'No data'}), 400

    events = data.get('events', [])
    
    event_counts = {}
    top_ips_src = defaultdict(int)
    top_ips_dst = defaultdict(int)
    port_dist = defaultdict(int)
    
    for event in events:
        event_type = event.get('type')
        event_counts[event_type] = event_counts.get(event_type, 0) + 1
        
        src = event.get('source_ip')
        dst = event.get('dest_ip')
        if src:
            top_ips_src[src] += 1
        if dst:
            top_ips_dst[dst] += 1
        
        details = event.get('details', {})
        dport = details.get('dport')
        if dport:
            port_dist[str(dport)] += 1

    return jsonify({
        'event_counts': event_counts,
        'top_sources': sorted(top_ips_src.items(), key=lambda x: x[1], reverse=True)[:10],
        'top_destinations': sorted(top_ips_dst.items(), key=lambda x: x[1], reverse=True)[:10],
        'port_distribution': dict(sorted(port_dist.items(), key=lambda x: x[1], reverse=True)[:15]),
        'total_events': len(events)
    })

@app.route('/api/threats')
def api_threats():
    from threat_analyzer import ThreatAnalyzer
    
    data = load_report_data()
    if not data:
        return jsonify({'error': 'No data'}), 400

    analyzer = ThreatAnalyzer(data.get('events', []), data.get('links', []))
    analyzer.detect_patterns()
    
    threat_scores = {}
    for event in data.get('events', []):
        threat_scores[event['id']] = analyzer.get_threat_score(event['id'])
    
    return jsonify({
        'summary': analyzer.get_threat_summary(),
        'threat_scores': threat_scores,
        'patterns': analyzer.attack_patterns
    })

@app.route('/api/search')
def api_search():
    query = request.args.get('q', '').lower()
    field = request.args.get('field', 'all')
    
    data = load_report_data()
    if not data:
        return jsonify({'error': 'No data'}), 400

    events = data.get('events', [])
    results = []

    for event in events:
        match = False
        if field == 'all':
            match = (query in str(event).lower())
        elif field == 'ip':
            match = (query in event.get('source_ip', '') or query in event.get('dest_ip', ''))
        elif field == 'domain':
            match = query in event.get('details', {}).get('query', '').lower() or query in event.get('details', {}).get('sni', '').lower()
        elif field == 'type':
            match = query in event.get('type', '').lower()
        
        if match:
            results.append({
                'id': event['id'],
                'type': event.get('type'),
                'timestamp': event.get('timestamp'),
                'source': event.get('source_ip'),
                'destination': event.get('dest_ip'),
                'description': event.get('description')
            })

    return jsonify({'results': results, 'count': len(results)})

@app.route('/api/geoip/<ip>')
def api_geoip(ip):
    from threat_analyzer import ThreatAnalyzer
    
    analyzer = ThreatAnalyzer([], [])
    geo_data = analyzer.analyze_geoip(ip)
    return jsonify(geo_data)

@app.route('/api/geoips')
def api_geoips():
    from threat_analyzer import ThreatAnalyzer
    
    data = load_report_data()
    if not data:
        return jsonify({'error': 'No data'}), 400
    
    analyzer = ThreatAnalyzer(data.get('events', []), data.get('links', []))
    
    all_ips = set()
    for event in data.get('events', []):
        src = event.get('source_ip')
        dst = event.get('dest_ip')
        if src:
            all_ips.add(src)
        if dst:
            all_ips.add(dst)
    
    geoips = []
    for ip in list(all_ips)[:50]:
        geo = analyzer.analyze_geoip(ip)
        geoips.append(geo)
    
    return jsonify({'locations': geoips})

def load_report_data():
    events_path = os.path.join(os.path.dirname(__file__), 'events.json')
    if not os.path.exists(events_path):
        return None
    try:
        with open(events_path, 'r') as f:
            return json.load(f)
    except Exception:
        return None

def format_timestamp(ts):
    try:
        return datetime.fromtimestamp(float(ts)).strftime('%Y-%m-%d %H:%M:%S')
    except Exception:
        return str(ts)

def summarize_event(event):
    event_type = event.get('type', '')
    details = event.get('details') or {}
    if event_type == 'HTTP Request':
        method = details.get('method') or ''
        uri = details.get('uri') or ''
        host = details.get('host') or ''
        return f"{method} {host}{uri}".strip()
    if event_type == 'HTTP Response':
        status = details.get('status') or ''
        reason = details.get('reason') or ''
        return f"{status} {reason}".strip()
    if event_type == 'DNS Query':
        return details.get('query') or details.get('qname') or ''
    if event_type == 'DNS Response':
        domain = details.get('domain') or details.get('name') or ''
        ip_addr = details.get('ip') or ''
        return f"{domain} -> {ip_addr}".strip()
    if event_type == 'TLS SNI':
        return details.get('sni') or ''
    if event_type == 'TCP Connection':
        sport = details.get('sport') or ''
        dport = details.get('dport') or ''
        return f"{sport} -> {dport}".strip()
    if event_type == 'ICMP':
        icmp_type = details.get('type')
        code = details.get('code')
        return f"type={icmp_type} code={code}".strip()
    if event_type == 'ARP':
        return details.get('operation') or ''
    return event.get('description') or ''

def build_report_rows(events):
    rows = []
    counts = {}
    for event in events:
        event_type = event.get('type', 'Unknown')
        counts[event_type] = counts.get(event_type, 0) + 1
        details = event.get('details') or {}
        rows.append({
            'time': format_timestamp(event.get('timestamp')),
            'type': event_type,
            'src': event.get('source_ip') or '',
            'dst': event.get('dest_ip') or '',
            'message': summarize_event(event),
            'details': json.dumps(details, indent=2, ensure_ascii=True)
        })
    return rows, counts

@app.route('/report/pdf')
def report_pdf():
    data = load_report_data()
    if not data or not data.get('events'):
        return jsonify({'error': 'No report data found. Analyze a PCAP first.'}), 400

    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    events = data['events']
    rows, counts = build_report_rows(events)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, title='Cyber Attack Storyteller Report')
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph('Cyber Attack Storyteller Report', styles['Title']))
    elements.append(Paragraph('Generated report of network events.', styles['Normal']))
    elements.append(Spacer(1, 12))

    summary_data = [['Event Type', 'Count']]
    for event_type, count in sorted(counts.items()):
        summary_data.append([event_type, str(count)])
    summary_table = Table(summary_data, colWidths=[300, 100])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
    ]))
    elements.append(Paragraph('Summary', styles['Heading2']))
    elements.append(summary_table)
    elements.append(Spacer(1, 12))

    event_table_data = [['Time', 'Type', 'Source', 'Destination', 'Message', 'Details']]
    for row in rows:
        event_table_data.append([
            row['time'],
            row['type'],
            row['src'],
            row['dst'],
            row['message'],
            row['details']
        ])

    event_table = Table(event_table_data, colWidths=[90, 70, 90, 90, 120, 200])
    event_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    elements.append(Paragraph('Event Details', styles['Heading2']))
    elements.append(event_table)

    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='cyber-attack-report.pdf', mimetype='application/pdf')

@app.route('/report/docx')
def report_docx():
    data = load_report_data()
    if not data or not data.get('events'):
        return jsonify({'error': 'No report data found. Analyze a PCAP first.'}), 400

    from docx import Document
    from docx.shared import Pt

    events = data['events']
    rows, counts = build_report_rows(events)

    document = Document()
    document.add_heading('Cyber Attack Storyteller Report', 0)
    document.add_paragraph('Generated report of network events.')

    document.add_heading('Summary', level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'Event Type'
    hdr_cells[1].text = 'Count'
    for event_type, count in sorted(counts.items()):
        row_cells = summary_table.add_row().cells
        row_cells[0].text = event_type
        row_cells[1].text = str(count)

    document.add_heading('Event Details', level=1)
    event_table = document.add_table(rows=1, cols=6)
    event_table.style = 'Table Grid'
    event_hdr = event_table.rows[0].cells
    event_hdr[0].text = 'Time'
    event_hdr[1].text = 'Type'
    event_hdr[2].text = 'Source'
    event_hdr[3].text = 'Destination'
    event_hdr[4].text = 'Message'
    event_hdr[5].text = 'Details'

    for row in rows:
        row_cells = event_table.add_row().cells
        row_cells[0].text = row['time']
        row_cells[1].text = row['type']
        row_cells[2].text = row['src']
        row_cells[3].text = row['dst']
        row_cells[4].text = row['message']
        row_cells[5].text = row['details']
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='cyber-attack-report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/upload', methods=['POST'])
def upload_file():
    if not SCAPY_OK:
        return jsonify({'error': 'Scapy is not installed on the server', 'details': 'Run: pip install scapy'}), 500

    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_name = f"{uuid.uuid4().hex}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
        file.save(filepath)
        print(f"[*] Saved uploaded file to {filepath}")

        try:
            parser_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pcap_parser.py')
            print(f"[*] Running parser: {sys.executable} {parser_script} {filepath}")
            
            # Use sys.executable to ensure we use the same Python environment
            result = subprocess.run(
                [sys.executable, parser_script, filepath],
                capture_output=True, text=True, cwd=os.path.dirname(os.path.abspath(__file__))
            )
            print(f"[*] Parser return code: {result.returncode}")
            if result.stdout:
                print("[*] Parser stdout:", result.stdout)
            if result.stderr:
                print("[!] Parser stderr:", result.stderr)

            if result.returncode != 0:
                error_msg = result.stderr or result.stdout or "Unknown error"
                return jsonify({'error': 'Parser failed', 'details': error_msg}), 500

            events_path = os.path.join(os.path.dirname(__file__), 'events.json')
            with open(events_path, 'r') as f:
                data = json.load(f)
            return jsonify(data)

        except Exception as e:
            print(f"[!] Exception in upload: {e}")
            return jsonify({'error': str(e)}), 500
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)
                print(f"[*] Deleted uploaded file {filepath}")
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/events.json')
def get_events():
    return send_from_directory(os.path.dirname(__file__), 'events.json')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
